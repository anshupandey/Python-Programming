# -*- coding: utf-8 -*-
"""
Created on Wed Dec 22 13:42:54 2021

@author: anshu
"""

import docx

# pip install python-docx


# create an instance of a word doc

doc = docx.Document()

# adding a heading
doc.add_heading("Python Programming Course")

para = doc.add_paragraph(""" This is a python course, which covers from basic to intermediate level.
                         It covers data types, control flow, operators, type conversion, functions, generators, lambda function and object oriented programming.
                         """)
                         
doc.save("Mydocument.docx")

##############################################


# reading data from an existing file

doc = docx.Document("Mydocument.docx")

for p in doc.paragraphs:
    print(p.text)
    
data = [[1,"Anshu",25,'Male'],
        [2,"Jessica",22,'Female'],
        [3,'James',32,'Male'],
        [4,'Johny',28,'Male']]


table = doc.add_table(rows=1,cols=4)

# adding heading to the table
row = table.rows[0].cells
row[0].text = "Serial Number"
row[1].text = "Name"
row[2].text = "Age"
row[3].text = "Gender"

for sr,name,age,gen in data:
    row = table.add_row().cells
    row[0].text = str(sr)
    row[1].text = name
    row[2].text = str(age)
    row[3].text = gen

table.style = "Colorful List"    
doc.save("Mydocument.docx")


##############################################

# Numbered list

doc.add_heading("Python Course ",0)
doc.add_heading("Python Course ",1)
doc.add_heading("Python Course ",2)

doc.add_heading("Style: List Number",2)
doc.add_paragraph("Python Fundamentals",style="List Number")
doc.add_paragraph("Python OOPs",style="List Number")
doc.add_paragraph("Python Packages",style="List Number")
doc.add_paragraph("Python RegeX",style="List Number")

doc.save("Mydocument.docx")
